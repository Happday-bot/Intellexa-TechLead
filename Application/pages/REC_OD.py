import streamlit as st
import pandas as pd
from Intellexa_Login import login
from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

# Redirect to login if not authenticated
if not st.session_state.get("authenticated", False):
    # login()
    # st.stop()
    # url = "https://intellexa-techlead.streamlit.app/"  # Replace with your desired URL
    # st.markdown(f'<meta http-equiv="refresh" content="0;url={url}">', unsafe_allow_html=True)
    st.stop()

def load_csv_file():
    st.info("Step 1: Please upload your CSV file.")
    uploaded_file = st.file_uploader("Upload CSV", type="csv")
    if uploaded_file is not None:
        try:
            df = pd.read_csv(uploaded_file)
            df.columns = df.columns.str.lower().str.strip()  # Convert headers to lowercase and strip whitespace for case insensitivity
            st.success("CSV successfully loaded!")
            return df
        except Exception as e:
            st.error(f"Error loading CSV file: {e}")
    return None

def check_null_values(df):
    if df.isnull().values.any():
        st.error("Missing values detected!")
        return True
    else:
        st.success("No null values detected!")
        return False

def process_data(df, event_type):
    required_columns = ['name', 'register number', 'college name', 'year of study', 'department', 'section']
    if event_type == "Inter College Event":
        required_columns.extend(['event name', 'email address'])
    missing_columns = [col for col in required_columns if col not in df.columns]
    if missing_columns:
        st.error(f"Missing columns: {missing_columns}")
        return None
    
    filtered_df = df[required_columns]
    if event_type == "Inter College Event":
        sorted_df = filtered_df.sort_values(by=['college name', 'event name'])
    else:
        sorted_df = filtered_df[filtered_df['college name'].str.contains("Rajalakshmi Engineering", na=False, case=False)]
        sorted_df = sorted_df.sort_values(by=['name', 'register number'])  # Ensure sorting for intra-college events
    
    return sorted_df

def create_word_document(df,image_path = None):
    doc = Document()
    
    # Add an image (if provided)
    if image_path:
        doc.add_picture(image_path, width=Inches(6))  # Adjust width as needed
        doc.add_paragraph("\n")  # Adds space after the image


     # Add sender and recipient details
    doc.add_paragraph("From\nIntellexa,\nRajalakshmi Engineering College, Thandalam, Chennai.\n")
    doc.add_paragraph("To\nHead of all Departments,\nRajalakshmi Engineering College, Thandalam, Chennai.\n")

    # Add subject and body text
    doc.add_paragraph("Respected Sir,")
    doc.add_paragraph("Sub: Request for On-Duty for {Data Analytics workshop and career guidance}")
    doc.add_paragraph(
        "We, the Intellexa club of Rajalakshmi Engineering College kindly request you to provide on-duty for our below-mentioned "
        "members for attending a {one-day} workshop organized in {collaboration with Chennai Data Circle on Data Analytics} "
        "in our college on {08-02-2025 (8:00 AM – 5:00 PM)}. Kindly consider our request and provide us the permission.\n"
    )

    # Add the closing statement and align to the right
    paragraph = doc.add_paragraph("\nYours Sincerely,\nIntellexa REC.")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Align everything to the right

    doc.add_paragraph("\n\n\n\n\n\n\n\n\n\n")

    # Add a table with the student data
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name.upper()

    # Add student data
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Save to a buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def attendance_sheet(df):
    # Create a Word document
    doc = Document()

    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Adjust margins (optional)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Add a table with headers
    table = doc.add_table(rows=1, cols=df.shape[1])
    table.style = "Table Grid"

    # Add column headers
    header_cells = table.rows[0].cells
    for j, column_name in enumerate(df.columns):
        header_cells[j].text = str(column_name)

    # Add data rows safely
    for i, row in df.iterrows():
        row_cells = table.add_row().cells  # Add new row safely
        for j, value in enumerate(row):
            row_cells[j].text = str(value)

    # Save to a BytesIO buffer
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer


def main():
    st.title("OD List Processor")
    st.markdown("Upload the Excel sheet as a CSV file.")
    
    event_type = st.radio("Select Event Type:", ["Inter College Event", "Intra College Event"])
    df = load_csv_file()
    
    if df is not None:
        st.markdown("### Data Snapshot")
        st.dataframe(df.head())
        
        if check_null_values(df):
            if st.radio("Proceed with null values?", ("Yes", "No")) == "No":
                st.experimental_rerun()  # Allow re-upload instead of stopping
        
        processed_df = process_data(df, event_type)
        if processed_df is not None:
            st.markdown("### Processed Data")
            st.dataframe(processed_df)
            
            csv_data = processed_df.to_csv(index=False).encode('utf-8')
            
            if event_type == "Inter College Event":
                rec_students = processed_df[processed_df['college name'].str.contains("Rajalakshmi Engineering", na=False, case=False)]
                outer_students = processed_df[~processed_df['college name'].str.contains("Rajalakshmi Engineering", na=False, case=False)]
                
                st.download_button("Download REC Students", rec_students.to_csv(index=False).encode('utf-8'), "rec_students.csv", "text/csv")
                st.download_button("Download Outer College Students", outer_students.to_csv(index=False).encode('utf-8'), "outer_students.csv", "text/csv")

                rec_students = rec_students.drop(columns=['email address'])
                # Convert CSV to Word
                docx_file = attendance_sheet(rec_students)

                # Download button
                st.download_button(
                label="Download Attendance sheet",
                data=docx_file,
                file_name="Attendance sheet.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                rec_students = rec_students.drop(columns=['college name','event name'])

                word_file = create_word_document(rec_students,"https://raw.githubusercontent.com/Happday-bot/Intellexa-TechLead/main/Application/assests/enhanced_logo.png")
                # image_url = "https://raw.githubusercontent.com/your-username/your-repo/main/logo.png"
                # word_file = create_word_document(rec_students, image_path=image_url)

                st.download_button("Download REC Students (Word)", word_file, "rec_students.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")



            else:
                rec_students = processed_df[processed_df['college name'].str.contains("Rajalakshmi Engineering", na=False, case=False)]
                rec_students = rec_students.drop(columns=['college name'])

                # Convert CSV to Word
                docx_file = attendance_sheet(rec_students)

                # Download button
                st.download_button(
                label="Download Attendance sheet",
                data=docx_file,
                file_name="Attendance sheet.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.download_button("Download Processed CSV", csv_data, "rec_students.csv", "text/csv")
                word_file = create_word_document(rec_students,"https://raw.githubusercontent.com/Happday-bot/Intellexa-TechLead/main/Application/assests/enhanced_logo.png")
                # image_url = "https://raw.githubusercontent.com/your-username/your-repo/main/logo.png"
                # word_file = create_word_document(rec_students, image_path=image_url)

                st.download_button("Download REC Students (Word)", word_file, "rec_students.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    main()

if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = True

if st.session_state["authenticated"]:
    if st.sidebar.button("🔓 Logout"):
        st.session_state["authenticated"] = False
        st.rerun()
